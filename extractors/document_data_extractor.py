"""
Document Data Extractor
Extracts project data from PDF, DOCX, and XLSX files to auto-fill forms
"""

import os
import json
import re
from io import BytesIO
from typing import Dict, Optional, Any

# PDF parsing
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# DOCX parsing
try:
    from docx import Document
except ImportError:
    Document = None

# XLSX parsing
try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


class DocumentDataExtractor:
    """Extract data from uploaded documents to fill forms"""
    
    # Field mappings for each document type - Spanish keywords to field names
    FIELD_MAPPINGS = {
        "estudios_previos": {
            "entidad": ["entidad", "entidad contratante", "institution"],
            "objeto": ["objeto", "objeto contractual", "objeto del contrato"],
            "presupuesto": ["presupuesto", "valor", "monto", "valor estimado"],
            "plazo": ["plazo", "duración", "tiempo de ejecución"],
            "modalidad": ["modalidad", "modalidad de selección"],
            "sector": ["sector", "área"],
            "descripcion": ["descripción", "justificación", "necesidad"],
        },
        "analisis_sector": {
            "sector": ["sector", "sector económico"],
            "entidad": ["entidad", "entidad contratante"],
            "objeto": ["objeto", "objeto a contratar"],
            "valor_estimado": ["valor", "presupuesto", "valor estimado"],
        },
        "dts": {
            "municipio": ["municipio", "ciudad"],
            "departamento": ["departamento"],
            "entidad": ["entidad"],
            "proyecto": ["proyecto", "nombre del proyecto"],
            "valor": ["valor", "presupuesto", "costo total"],
        },
        "certificaciones": {
            "nombre_proyecto": ["proyecto", "nombre del proyecto"],
            "municipio": ["municipio"],
            "departamento": ["departamento"],
            "valor": ["valor", "presupuesto"],
            "responsable": ["responsable", "formulador", "secretario"],
        },
        "mga_subsidios": {
            "municipio": ["municipio", "ciudad"],
            "departamento": ["departamento"],
            "entidad": ["entidad", "alcaldía"],
            "bpin": ["bpin", "código bpin"],
            "nombre_proyecto": ["proyecto", "nombre del proyecto", "título"],
            "valor_total": ["valor", "valor total", "presupuesto", "monto"],
            "duracion": ["duración", "plazo", "días"],
            "responsable": ["responsable", "formulador", "secretario"],
            "cargo": ["cargo", "puesto"],
            "plan_nacional": ["plan nacional", "pnd"],
            "plan_departamental": ["plan departamental"],
            "plan_municipal": ["plan municipal", "pdm"],
        }
    }
    
    def __init__(self, llm=None):
        """Initialize extractor with optional LLM for AI extraction"""
        self.llm = llm
    
    def extract_from_file(self, file, file_type: str, doc_type: str, user_context: str = "") -> Dict[str, Any]:
        """
        Extract data from uploaded file
        
        Args:
            file: Uploaded file object (Streamlit UploadedFile)
            file_type: Extension (.pdf, .docx, .xlsx)
            doc_type: Document type (estudios_previos, mga_subsidios, etc.)
            user_context: Optional user-provided context for extraction
            
        Returns:
            Dictionary with extracted field values
        """
        # Read file content
        if hasattr(file, 'read'):
            content = file.read()
            file.seek(0)  # Reset for potential re-read
        else:
            with open(file, 'rb') as f:
                content = f.read()
        
        # Extract text based on file type
        if file_type.lower() in ['.pdf', 'pdf']:
            text = self._extract_pdf_text(content)
        elif file_type.lower() in ['.docx', 'docx']:
            text = self._extract_docx_text(content)
        elif file_type.lower() in ['.xlsx', 'xlsx', '.xls', 'xls']:
            text = self._extract_xlsx_text(content)
        else:
            return {"error": f"Unsupported file type: {file_type}"}
        
        
        # Extract structured data
        if self.llm:
            result = self._extract_with_ai(text, doc_type, user_context)
        else:
            result = self._extract_with_patterns(text, doc_type)
            
        # Add raw text dump for context fallback
        result["context_dump"] = text
        
        # Store user context if provided
        if user_context:
            result["user_context"] = user_context
        
        return result
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        text_parts = []
        
        # Try PyMuPDF first (faster)
        if fitz:
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return "\n".join(text_parts)
            except Exception as e:
                print(f"PyMuPDF error: {e}")
        
        # Fall back to pdfplumber
        if pdfplumber:
            try:
                with pdfplumber.open(BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                return "\n".join(text_parts)
            except Exception as e:
                print(f"pdfplumber error: {e}")
        
        return ""
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        if not Document:
            return ""
        
        try:
            doc = Document(BytesIO(content))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""
    
    def _extract_xlsx_text(self, content: bytes) -> str:
        """Extract text from XLSX"""
        text_parts = []
        
        if pd:
            try:
                # Read all sheets
                xlsx = pd.ExcelFile(BytesIO(content))
                for sheet_name in xlsx.sheet_names:
                    df = pd.read_excel(xlsx, sheet_name=sheet_name)
                    # Convert to text format
                    for col in df.columns:
                        text_parts.append(f"{col}: {df[col].tolist()}")
                return "\n".join(text_parts)
            except Exception as e:
                print(f"Pandas Excel error: {e}")
        
        if openpyxl:
            try:
                wb = openpyxl.load_workbook(BytesIO(content))
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        row_text = [str(cell) for cell in row if cell]
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                return "\n".join(text_parts)
            except Exception as e:
                print(f"openpyxl error: {e}")
        
        return ""
    
    def _extract_with_patterns(self, text: str, doc_type: str) -> Dict[str, str]:
        """Extract data using regex patterns and keyword matching"""
        result = {}
        mappings = self.FIELD_MAPPINGS.get(doc_type, {})
        
        lines = text.split('\n')
        text_lower = text.lower()
        
        for field_name, keywords in mappings.items():
            for keyword in keywords:
                # Try to find keyword followed by colon and value
                pattern = rf'{keyword}\s*[:\-]?\s*(.+?)(?:\n|$)'
                match = re.search(pattern, text_lower, re.IGNORECASE)
                if match:
                    value = match.group(1).strip()
                    # Clean up the value
                    value = re.sub(r'^[:\-\s]+', '', value)
                    if value and len(value) > 1:
                        result[field_name] = value[:500]  # Limit length
                        break
        
        return result
    
    def _extract_with_ai(self, text: str, doc_type: str, user_context: str = "") -> Dict[str, str]:
        """Extract data using AI/LLM with improved prompt for better quality"""
        from langchain_core.prompts import ChatPromptTemplate
        from langchain_core.output_parsers import StrOutputParser
        
        # Comprehensive field list for all document types (unified)
        all_fields = [
            "municipio", "departamento", "entidad", "bpin", "nombre_proyecto",
            "valor_total", "duracion", "responsable", "cargo", "alcalde",
            "objeto", "necesidad", "alcance", "modalidad", "fuente_financiacion",
            "sector", "codigo_ciiu", "codigos_unspsc", "programa", "subprograma",
            "plan_nacional", "plan_departamental", "plan_municipal",
            "poblacion_beneficiada", "indicador_producto", "meta_producto",
            "es_actualizacion"
        ]
        
        fields_str = ", ".join(all_fields)
        
        # Use MORE text for better extraction (increased from 6000 to 15000)
        # Also sample from different parts of the document to get data from all pages
        text_length = len(text)
        if text_length > 15000:
            # Smart sampling: take from beginning, middle, and end
            beginning = text[:6000]
            middle_start = max(0, (text_length // 2) - 3000)
            middle = text[middle_start:middle_start + 6000]
            end = text[-3000:]
            text_to_analyze = f"{beginning}\n\n[...SECCIÓN INTERMEDIA...]\n\n{middle}\n\n[...SECCIÓN FINAL...]\n\n{end}"
        else:
            text_to_analyze = text
        
        # Build context section if provided
        context_section = ""
        if user_context:
            context_section = """

===== CONTEXTO DEL USUARIO =====
""" + user_context + """
===== FIN DEL CONTEXTO =====

NOTA: El usuario ha proporcionado contexto adicional. Si indica que es una ACTUALIZACIÓN, 
agrega "es_actualizacion": "Si" al JSON. Prioriza la información según las instrucciones del usuario.
"""
        
        # Build the human message without f-string to avoid brace escaping issues
        human_message = """Extrae los siguientes campos del documento gubernamental colombiano:
""" + fields_str + context_section + """

===== DOCUMENTO COMPLETO =====
""" + text_to_analyze + """
===== FIN DEL DOCUMENTO =====

IMPORTANTE: Analiza tablas, encabezados, y texto en cualquier formato.
Responde ÚNICAMENTE con un JSON válido. Ejemplo de formato:
municipio: San Pablo
departamento: Bolívar
entidad: Alcaldía Municipal
bpin: 2024000001234
nombre_proyecto: Construcción de acueducto
valor_total: 500000000
duracion: 180
responsable: Juan Pérez
cargo: Secretario de Planeación
objeto: Contratar la construcción...
necesidad: Se requiere mejorar...
sector: Agua Potable
es_actualizacion: No

Responde con JSON usando las claves anteriores."""

        prompt = ChatPromptTemplate.from_messages([
            ("system", """Eres un experto en extracción de datos de documentos gubernamentales colombianos.
Tu especialidad es extraer información de:
- Metodología General Ajustada (MGA)
- Estudios Previos
- Convenios Interadministrativos
- Certificaciones municipales
- Documentos Técnicos de Soporte (DTS)

INSTRUCCIONES CRÍTICAS:
1. Lee TODO el documento cuidadosamente
2. Extrae TODOS los valores que encuentres, incluso si están en tablas o listas
3. Para "valor_total" busca: presupuesto, valor del contrato, monto, costo total
4. Para "objeto" busca: objeto contractual, objeto del convenio, descripción del proyecto
5. Para "necesidad" busca: justificación, problema a resolver, antecedentes
6. Para "alcance" busca: actividades, productos, entregables
7. Para "municipio" y "departamento" busca: ubicación, localización geográfica
8. Para "responsable" busca: secretario, formulador, representante legal
9. Si no encuentras un campo exacto, busca sinónimos o información relacionada
10. NO inventes datos. Si no está, omítelo del JSON.
11. Responde SOLO con JSON válido, nada más."""),
            ("human", human_message)
        ])
        
        try:
            chain = prompt | self.llm | StrOutputParser()
            response = chain.invoke({})
            
            # Parse JSON response - try multiple patterns
            import json
            
            # Try to find JSON in response
            json_patterns = [
                r'\{[\s\S]*\}',  # Any JSON object
                r'```json\s*([\s\S]*?)\s*```',  # Code block
                r'```\s*([\s\S]*?)\s*```',  # Generic code block
            ]
            
            for pattern in json_patterns:
                match = re.search(pattern, response, re.DOTALL)
                if match:
                    try:
                        json_str = match.group(1) if '```' in pattern else match.group(0)
                        result = json.loads(json_str)
                        # Filter out null values and empty strings
                        return {k: v for k, v in result.items() if v and v != "null" and str(v).strip()}
                    except json.JSONDecodeError:
                        continue
                        
        except Exception as e:
            print(f"AI extraction error: {e}")
        
        # Fall back to pattern extraction
        return self._extract_with_patterns(text, doc_type)


def extract_data_from_upload(uploaded_file, doc_type: str, llm=None, user_context: str = "") -> Dict[str, Any]:
    """
    Convenience function to extract data from an uploaded file
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        doc_type: Type of document to generate
        llm: Optional LLM for AI-powered extraction
        user_context: Optional user-provided context (e.g., "this is for updating existing document")
        
    Returns:
        Dictionary with extracted field values
    """
    if not uploaded_file:
        return {}
    
    # Get file extension
    filename = uploaded_file.name if hasattr(uploaded_file, 'name') else str(uploaded_file)
    ext = os.path.splitext(filename)[1].lower()
    
    extractor = DocumentDataExtractor(llm=llm)
    return extractor.extract_from_file(uploaded_file, ext, doc_type, user_context=user_context)
