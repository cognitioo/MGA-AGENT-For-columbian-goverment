"""
MGA Subsidios Document Builder
Generates Word documents for complete MGA Subsidios (24 pages)
"""

import os
import re
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MGASubsidiosBuilder:
    """Builder for MGA Subsidios Word documents (24 pages)"""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self.doc = None
    
    def _safe_list(self, value, default=None):
        """Safely convert a value to a list, returning default if not iterable"""
        if default is None:
            default = []
        if isinstance(value, list):
            return value
        return default
    
    def build(self, data: dict, ai_content: dict, letterhead_file=None) -> str:
        """Build the MGA Subsidios document"""
        self._has_letterhead = letterhead_file is not None
        
        if letterhead_file:
            self.doc = self._load_template(letterhead_file)
        else:
            self.doc = Document()
        
        self._apply_styles()
        
        # Page 1 - Datos Básicos
        self._add_page_1_datos_basicos(ai_content.get("pagina_1_datos_basicos", {}))
        
        self.doc.add_page_break()
        
        # Page 2 - Plan de Desarrollo
        self._add_page_2_plan_desarrollo(ai_content.get("pagina_2_plan_desarrollo", {}))
        
        self.doc.add_page_break()
        
        # Page 3 - Problemática
        self._add_page_3_problematica(ai_content.get("pagina_3_problematica", {}))
        
        self.doc.add_page_break()
        
        # Page 4 - Causas y Efectos
        self._add_page_4_causas_efectos(ai_content.get("pagina_4_causas_efectos", {}))
        
        self.doc.add_page_break()
        
        # Page 5 - Participantes
        self._add_page_5_participantes(ai_content.get("pagina_5_participantes", {}))
        
        self.doc.add_page_break()
        
        # Pages 6-11
        self._add_pages_6_11(ai_content)
        
        self.doc.add_page_break()
        
        # Pages 12-16
        self._add_pages_12_16(ai_content)
        
        self.doc.add_page_break()
        
        # Pages 17-21
        self._add_pages_17_21(ai_content)
        
        self.doc.add_page_break()
        
        # Pages 22-24
        self._add_pages_22_24(ai_content)
        
        # Save and return
        return self._save_document(data)
    
    def _load_template(self, letterhead_file):
        """Load template from uploaded file"""
        try:
            if hasattr(letterhead_file, 'read'):
                return Document(BytesIO(letterhead_file.read()))
            else:
                return Document(letterhead_file)
        except Exception:
            return Document()
    
    def _apply_styles(self):
        """Apply document styles"""
        for section in self.doc.sections:
            if not self._has_letterhead:
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
    
    def _add_header(self, section_name: str, subsection: str = ""):
        """Add standard MGA header"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if subsection:
            run = p.add_run(f"{section_name} / ")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 128, 128)  # Teal
            run2 = p.add_run(subsection)
            run2.font.size = Pt(10)
            run2.font.color.rgb = RGBColor(0, 128, 128)
            run2.bold = True
        else:
            run = p.add_run(section_name)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 128, 128)
            run.bold = True
    
    def _add_section_title(self, title: str, color_hex: str = "0099CC"):
        """Add section title in blue"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 153, 204)  # Blue
        run.bold = True
        p.paragraph_format.space_after = Pt(12)
    
    def _add_subsection_title(self, title: str):
        """Add subsection title"""
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 128, 128)  # Teal
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    
    def _add_field(self, label: str, value: str):
        """Add a label-value field"""
        p = self.doc.add_paragraph()
        run_label = p.add_run(label)
        run_label.font.size = Pt(10)
        run_label.font.color.rgb = RGBColor(0, 128, 128)
        run_value = p.add_run(f"\n{value}")
        run_value.font.size = Pt(10)
    
    def _add_page_1_datos_basicos(self, content: dict):
        """Add Page 1 - Datos Básicos"""
        self._add_header("Datos básicos")
        
        # Title section (dark blue bar would be here in actual MGA)
        p_title = self.doc.add_paragraph()
        run = p_title.add_run(content.get("titulo_documento", ""))
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        self._add_section_title("Datos básicos")
        self._add_subsection_title("01 - Datos básicos del proyecto")
        
        self._add_field("Nombre", content.get("nombre", ""))
        
        # Two columns: Tipología and Código BPIN
        p = self.doc.add_paragraph()
        run1 = p.add_run("Tipología\n")
        run1.font.size = Pt(10)
        run1.font.color.rgb = RGBColor(0, 128, 128)
        run1v = p.add_run(content.get("tipologia", ""))
        run1v.font.size = Pt(10)
        
        p.add_run("                    ")
        
        run2 = p.add_run("Código BPIN\n")
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0, 128, 128)
        run2v = p.add_run(content.get("codigo_bpin", ""))
        run2v.font.size = Pt(10)
        
        self._add_field("Sector", content.get("sector", ""))
        
        # Es Proyecto Tipo and Fecha creación
        p2 = self.doc.add_paragraph()
        p2.add_run("Es Proyecto Tipo: ").font.color.rgb = RGBColor(0, 128, 128)
        p2.add_run(content.get("es_proyecto_tipo", "No"))
        p2.add_run("                    ")
        p2.add_run("Fecha creación: ").font.color.rgb = RGBColor(0, 128, 128)
        p2.add_run(content.get("fecha_creacion", ""))
        
        self._add_field("Identificador:", content.get("identificador", ""))
        self._add_field("Formulador Ciudadano:", content.get("formulador_ciudadano", ""))
        self._add_field("Formulador Oficial:", content.get("formulador_oficial", ""))
    
    def _add_page_2_plan_desarrollo(self, content: dict):
        """Add Page 2 - Plan de Desarrollo"""
        self._add_header("Identificación", "Plan de desarrollo")
        
        self._add_section_title("Contribución a la política pública")
        
        # Plan Nacional
        self._add_subsection_title("01 - Contribución al Plan Nacional de Desarrollo")
        
        plan_nacional = content.get("plan_nacional", {})
        self._add_field("Plan", plan_nacional.get("nombre", ""))
        self._add_field("Programa", plan_nacional.get("programa", ""))
        
        # Table
        table = self.doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        
        headers = ["Transformación", "Pilar", "Catalizador", "Componente"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._set_cell_shading(cell, "0099CC")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
        
        table.rows[1].cells[0].text = plan_nacional.get("transformacion", "")
        table.rows[1].cells[1].text = plan_nacional.get("pilar", "")
        table.rows[1].cells[2].text = plan_nacional.get("catalizador", "")
        table.rows[1].cells[3].text = plan_nacional.get("componente", "")
        
        self.doc.add_paragraph()
        
        # Plan Departamental
        self._add_subsection_title("02 - Plan de Desarrollo Departamental o Sectorial")
        plan_depto = content.get("plan_departamental", {})
        self._add_field("Plan de Desarrollo Departamental o Sectorial", plan_depto.get("nombre", ""))
        self._add_field("Estrategia del Plan de Desarrollo Departamental o Sectorial", plan_depto.get("estrategia", ""))
        self._add_field("Programa del Plan Desarrollo Departamental o Sectorial", plan_depto.get("programa", ""))
        
        # Plan Municipal
        self._add_subsection_title("03 - Plan de Desarrollo Distrital o Municipal")
        plan_mun = content.get("plan_municipal", {})
        self._add_field("Plan de Desarrollo Distrital o Municipal", plan_mun.get("nombre", ""))
        self._add_field("Estrategia del Plan de Desarrollo Distrital o Municipal", plan_mun.get("estrategia", ""))
        self._add_field("Programa del Plan desarrollo Distrital o Municipal", plan_mun.get("programa", ""))
        
        # Grupos étnicos
        self._add_subsection_title("04 - Instrumentos de planeación de grupos étnicos")
        self._add_field("Tipo de entidad", content.get("instrumentos_grupos_etnicos", "No aplica"))
    
    def _add_page_3_problematica(self, content: dict):
        """Add Page 3 - Problemática"""
        self._add_header("Identificación", "Problemática")
        
        self._add_section_title("Identificación y descripción del problema")
        
        self._add_subsection_title("Problema central")
        p = self.doc.add_paragraph()
        p.add_run(content.get("problema_central", ""))
        
        self._add_subsection_title("Descripción de la situación existente con respecto al problema")
        desc = content.get("descripcion_situacion", "").replace("\\n", "\n")
        for para in desc.split("\n\n"):
            p = self.doc.add_paragraph()
            p.add_run(para)
            p.paragraph_format.space_after = Pt(6)
        
        self._add_subsection_title("Magnitud actual del problema – indicadores de referencia")
        mag = content.get("magnitud_problema", "").replace("\\n", "\n")
        for para in mag.split("\n\n"):
            p = self.doc.add_paragraph()
            p.add_run(para)
            p.paragraph_format.space_after = Pt(6)
    
    def _add_page_4_causas_efectos(self, content: dict):
        """Add Page 4 - Causas y Efectos"""
        self._add_header("Identificación", "Problemática")
        
        # Causas
        self._add_subsection_title("01 - Causas que generan el problema")
        
        causas_table = self.doc.add_table(rows=1, cols=2)
        causas_table.style = 'Table Grid'
        
        causas_table.rows[0].cells[0].text = "Causas directas"
        causas_table.rows[0].cells[1].text = "Causas indirectas"
        self._set_cell_shading(causas_table.rows[0].cells[0], "0099CC")
        self._set_cell_shading(causas_table.rows[0].cells[1], "0099CC")
        for cell in causas_table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        causas_directas = self._safe_list(content.get("causas_directas", []))
        causas_indirectas = self._safe_list(content.get("causas_indirectas", []))
        
        max_rows = max(len(causas_directas), len(causas_indirectas))
        for i in range(max_rows):
            row = causas_table.add_row()
            if i < len(causas_directas):
                c = causas_directas[i]
                row.cells[0].text = f"{c.get('numero', '')}. {c.get('causa', '')}"
            if i < len(causas_indirectas):
                c = causas_indirectas[i]
                row.cells[1].text = f"{c.get('numero', '')} {c.get('causa', '')}"
        
        self.doc.add_paragraph()
        
        # Efectos
        self._add_subsection_title("02 - Efectos generados por el problema")
        
        efectos_table = self.doc.add_table(rows=1, cols=2)
        efectos_table.style = 'Table Grid'
        
        efectos_table.rows[0].cells[0].text = "Efectos directos"
        efectos_table.rows[0].cells[1].text = "Efectos indirectos"
        self._set_cell_shading(efectos_table.rows[0].cells[0], "0099CC")
        self._set_cell_shading(efectos_table.rows[0].cells[1], "0099CC")
        for cell in efectos_table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        efectos_directos = self._safe_list(content.get("efectos_directos", []))
        efectos_indirectos = self._safe_list(content.get("efectos_indirectos", []))
        
        max_rows = max(len(efectos_directos), len(efectos_indirectos))
        for i in range(max_rows):
            row = efectos_table.add_row()
            if i < len(efectos_directos):
                e = efectos_directos[i]
                row.cells[0].text = f"{e.get('numero', '')}. {e.get('efecto', '')}"
            if i < len(efectos_indirectos):
                e = efectos_indirectos[i]
                row.cells[1].text = f"{e.get('numero', '')} {e.get('efecto', '')}"
    
    def _add_page_5_participantes(self, content: dict):
        """Add Page 5 - Participantes"""
        self._add_header("Identificación", "Participantes")
        
        self._add_section_title("Identificación y análisis de participantes")
        
        self._add_subsection_title("01 - Identificación de los participantes")
        
        # Participants table
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = "Participante"
        table.rows[0].cells[1].text = "Contribución o Gestión"
        self._set_cell_shading(table.rows[0].cells[0], "0099CC")
        self._set_cell_shading(table.rows[0].cells[1], "0099CC")
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        for part in self._safe_list(content.get("participantes", [])):
            row = table.add_row()
            
            # Left cell - participant details
            cell_text = f"Actor: {part.get('actor', '')}\n"
            cell_text += f"Entidad: {part.get('entidad', '')}\n"
            cell_text += f"Posición: {part.get('posicion', '')}\n"
            cell_text += f"Intereses o Expectativas: {part.get('intereses', '')}"
            row.cells[0].text = cell_text
            
            # Right cell - contribution
            row.cells[1].text = part.get("contribucion", "")
        
        self.doc.add_paragraph()
        
        # Análisis
        self._add_subsection_title("02 - Análisis de los participantes")
        analisis = content.get("analisis_participantes", "").replace("\\n", "\n")
        for para in analisis.split("\n\n"):
            p = self.doc.add_paragraph()
            p.add_run(para)
            p.paragraph_format.space_after = Pt(6)
    
    def _add_page_6_poblacion(self, content: dict):
        """Add Page 6 - Población"""
        self._add_header("Identificación", "Población")
        
        # Población afectada
        self._add_subsection_title("01 - Población afectada por el problema")
        
        pob_afectada = content.get("poblacion_afectada", {})
        self._add_field("Tipo de población", pob_afectada.get("tipo", "Personas"))
        self._add_field("Número", pob_afectada.get("numero", ""))
        self._add_field("Fuente de la información", pob_afectada.get("fuente", ""))
        
        # Localización table
        self._add_subsection_title("Localización")
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = "Ubicación general"
        table.rows[0].cells[1].text = "Localización específica/Otro tipo de entidad étnica"
        self._set_cell_shading(table.rows[0].cells[0], "0099CC")
        self._set_cell_shading(table.rows[0].cells[1], "0099CC")
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        table.rows[1].cells[0].text = f"Región: {pob_afectada.get('region', 'Caribe')}"
        table.rows[2].cells[0].text = f"Departamento: {pob_afectada.get('departamento', '')}"
        table.rows[3].cells[0].text = f"Municipio: {pob_afectada.get('municipio', '')}"
        table.rows[4].cells[0].text = "Tipo de Agrupación:\nAgrupación:"
        
        self.doc.add_paragraph()
        
        # Población objetivo
        self._add_subsection_title("02 - Población objetivo de la intervención")
        
        pob_objetivo = content.get("poblacion_objetivo", {})
        self._add_field("Tipo de población", pob_objetivo.get("tipo", "Personas"))
        self._add_field("Número", pob_objetivo.get("numero", ""))
        self._add_field("Fuente de la información", pob_objetivo.get("fuente", ""))
        
        # Localización table for objetivo
        self._add_subsection_title("Localización")
        table2 = self.doc.add_table(rows=5, cols=3)
        table2.style = 'Table Grid'
        
        table2.rows[0].cells[0].text = "Ubicación general"
        table2.rows[0].cells[1].text = "Localización específica/Otro tipo de entidad étnica"
        table2.rows[0].cells[2].text = "Nombre del consejo comunitario"
        for i in range(3):
            self._set_cell_shading(table2.rows[0].cells[i], "0099CC")
            for run in table2.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        table2.rows[1].cells[0].text = f"Región: {pob_objetivo.get('region', 'Caribe')}"
        table2.rows[2].cells[0].text = f"Departamento: {pob_objetivo.get('departamento', '')}"
        table2.rows[3].cells[0].text = f"Municipio: {pob_objetivo.get('municipio', '')}"
        table2.rows[4].cells[0].text = "Tipo de Agrupación:\nAgrupación:"
    
    def _add_page_7_objetivos(self, content: dict):
        """Add Page 7 - Objetivos"""
        self._add_header("Identificación", "Objetivos")
        
        self._add_section_title("Objetivos específicos")
        
        self._add_subsection_title("01 - Objetivo general e indicadores de seguimiento")
        
        self._add_field("Problema central", content.get("problema_central", ""))
        self._add_field("Objetivo general – Propósito", content.get("objetivo_general", ""))
        
        # Indicadores table
        self._add_subsection_title("Indicadores para medir el objetivo general")
        
        indicadores = content.get("indicadores", [])
        if indicadores:
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            table.rows[0].cells[0].text = "Indicador objetivo"
            table.rows[0].cells[1].text = "Descripción"
            table.rows[0].cells[2].text = "Fuente de verificación"
            for cell in table.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for ind in indicadores:
                row = table.add_row()
                row.cells[0].text = ind.get("nombre", "")
                row.cells[1].text = f"Medido a través de: {ind.get('medido', '')}\nMeta: {ind.get('meta', '')}\nTipo de fuente: {ind.get('tipo_fuente', '')}"
                row.cells[2].text = ind.get("fuente_verificacion", "")
        
        self.doc.add_paragraph()
        
        # Relación causas-objetivos
        self._add_subsection_title("02 - Relaciones entre las causas y objetivos")
        
        relaciones = content.get("relacion_causas_objetivos", [])
        if relaciones:
            table2 = self.doc.add_table(rows=1, cols=2)
            table2.style = 'Table Grid'
            
            table2.rows[0].cells[0].text = "Causa relacionada"
            table2.rows[0].cells[1].text = "Objetivos específicos"
            for cell in table2.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for rel in relaciones:
                row = table2.add_row()
                row.cells[0].text = rel.get("causa", "")
                row.cells[1].text = rel.get("objetivo", "")
        
        self.doc.add_paragraph()
        
        # Alternativas
        self._add_section_title("Alternativas de la solución")
        self._add_subsection_title("01 - Alternativas de la solución")
        
        alternativas = content.get("alternativas", [])
        if alternativas:
            table3 = self.doc.add_table(rows=1, cols=3)
            table3.style = 'Table Grid'
            
            table3.rows[0].cells[0].text = "Nombre de la alternativa"
            table3.rows[0].cells[1].text = "Se evaluará con esta herramienta"
            table3.rows[0].cells[2].text = "Estado"
            for cell in table3.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for alt in alternativas:
                row = table3.add_row()
                row.cells[0].text = alt.get("nombre", "")
                row.cells[1].text = alt.get("evaluacion", "")
                row.cells[2].text = alt.get("estado", "")
        
        # Evaluaciones
        self._add_subsection_title("Evaluaciones a realizar")
        evaluaciones = content.get("evaluaciones", {})
        
        table4 = self.doc.add_table(rows=3, cols=2)
        table4.style = 'Table Grid'
        table4.rows[0].cells[0].text = "Rentabilidad:"
        table4.rows[0].cells[1].text = evaluaciones.get("rentabilidad", "Si")
        table4.rows[1].cells[0].text = "Costo - Eficiencia y Costo mínimo:"
        table4.rows[1].cells[1].text = evaluaciones.get("costo_eficiencia", "Si")
        table4.rows[2].cells[0].text = "Evaluación multicriterio:"
        table4.rows[2].cells[1].text = evaluaciones.get("multicriterio", "No")
    
    def _add_estudio_necesidades_servicio(self, servicio_data: dict, servicio_nombre: str):
        """Add estudio de necesidades for a service"""
        self._add_header("Identificación", "Objetivos")
        
        p = self.doc.add_paragraph()
        run = p.add_run(f"Alternativa 1. Transferir los recursos de subsidios para los servicios públicos domiciliarios de Acueducto, Alcantarillado y Aseo a la EMACALA S.A.S E.S.P")
        run.bold = True
        run.font.size = Pt(10)
        
        self._add_section_title("Estudio de necesidades")
        self._add_subsection_title("01 - Bien o servicio")
        
        self._add_field("Bien o servicio", servicio_data.get("bien_servicio", ""))
        self._add_field("Medido a través de", servicio_data.get("medido", "Número"))
        self._add_field("Descripción", servicio_data.get("descripcion", ""))
        self._add_field("Descripción de la Demanda", servicio_data.get("descripcion_demanda", ""))
        self._add_field("Descripción de la Oferta", servicio_data.get("descripcion_oferta", ""))
        
        # Oferta/Demanda table
        tabla_od = servicio_data.get("tabla_oferta_demanda", [])
        if tabla_od:
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            headers = ["Año", "Oferta", "Demanda", "Déficit"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                self._set_cell_shading(table.rows[0].cells[i], "0099CC")
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for item in tabla_od:
                row = table.add_row()
                row.cells[0].text = str(item.get("ano", ""))
                row.cells[1].text = str(item.get("oferta", ""))
                row.cells[2].text = str(item.get("demanda", ""))
                row.cells[3].text = str(item.get("deficit", ""))
                
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _add_pages_6_11(self, content: dict):
        """Add pages 6-11"""
        # Page 6 - Población
        self._add_page_6_poblacion(content.get("pagina_6_poblacion", {}))
        
        self.doc.add_page_break()
        
        # Page 7 - Objetivos
        self._add_page_7_objetivos(content.get("pagina_7_objetivos", {}))
        
        self.doc.add_page_break()
        
        # Pages 8-11 - Estudio de Necesidades
        estudio = content.get("pagina_8_9_10_11_estudio_necesidades", {})
        
        # Aseo
        self._add_estudio_necesidades_servicio(estudio.get("aseo", {}), "Aseo")
        
        self.doc.add_page_break()
        
        # Alcantarillado
        self._add_estudio_necesidades_servicio(estudio.get("alcantarillado", {}), "Alcantarillado")
        
        self.doc.add_page_break()
        
        # Acueducto
        self._add_estudio_necesidades_servicio(estudio.get("acueducto", {}), "Acueducto")
    
    def _add_page_12_analisis_tecnico(self, content: dict):
        """Add Page 12 - Análisis Técnico"""
        self._add_header("Preparación", "Análisis técnico")
        
        p = self.doc.add_paragraph()
        run = p.add_run("Alternativa: Transferir los recursos de subsidios para los servicios públicos domiciliarios de Acueducto, Alcantarillado y Aseo a la EMACALA S.A.S E.S.P")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        self._add_section_title("Análisis técnico de la alternativa")
        self._add_subsection_title("01 - Análisis técnico de la alternativa")
        
        self._add_subsection_title("Análisis técnico de la alternativa")
        p = self.doc.add_paragraph()
        p.add_run(content.get("descripcion_alternativa", ""))
        
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run(content.get("descripcion_subsidios", ""))
        
        # Datos de línea base
        datos = content.get("datos_linea_base", {})
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run(f"Con corte a {datos.get('fecha_corte', 'diciembre de 2025')}, se presentan los siguientes datos de linea base:")
        
        # Acueducto
        p = self.doc.add_paragraph()
        run = p.add_run("Servicio de Acueducto")
        run.bold = True
        
        acueducto = datos.get("acueducto", {})
        e1 = acueducto.get("estrato_1", {})
        e2 = acueducto.get("estrato_2", {})
        p = self.doc.add_paragraph()
        p.add_run(f"  Estrato 1: {e1.get('usuarios', '')} usuarios — Tarifa plena: {e1.get('tarifa_plena', '')} — Subsidio: {e1.get('subsidio', '')}")
        p = self.doc.add_paragraph()
        p.add_run(f"  Estrato 2: {e2.get('usuarios', '')} usuarios — Tarifa plena: {e2.get('tarifa_plena', '')} — Subsidio: {e2.get('subsidio', '')}")
        
        # Alcantarillado
        p = self.doc.add_paragraph()
        run = p.add_run("Servicio de Alcantarillado")
        run.bold = True
        
        alcant = datos.get("alcantarillado", {})
        e1 = alcant.get("estrato_1", {})
        e2 = alcant.get("estrato_2", {})
        p = self.doc.add_paragraph()
        p.add_run(f"  Estrato 1: {e1.get('usuarios', '')} usuarios — Tarifa plena: {e1.get('tarifa_plena', '')} — Subsidio: {e1.get('subsidio', '')}")
        p = self.doc.add_paragraph()
        p.add_run(f"  Estrato 2: {e2.get('usuarios', '')} usuarios — Tarifa plena: {e2.get('tarifa_plena', '')} — Subsidio: {e2.get('subsidio', '')}")
        
        # Aseo
        p = self.doc.add_paragraph()
        run = p.add_run("Servicio de Aseo")
        run.bold = True
        
        aseo = datos.get("aseo", {})
        e1 = aseo.get("estrato_1", {})
        e2 = aseo.get("estrato_2", {})
        p = self.doc.add_paragraph()
        p.add_run(f"  Estrato 1: {e1.get('usuarios', '')} usuarios — Tarifa plena: {e1.get('tarifa_plena', '')} — Subsidio: {e1.get('subsidio', '')}")
        p = self.doc.add_paragraph()
        p.add_run(f"  Estrato 2: {e2.get('usuarios', '')} usuarios — Tarifa plena: {e2.get('tarifa_plena', '')} — Subsidio: {e2.get('subsidio', '')}")
        
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run(content.get("implementacion", ""))
        
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run(content.get("beneficio", ""))
    
    def _add_page_13_localizacion(self, content: dict):
        """Add Page 13 - Localización"""
        self._add_header("Preparación", "Localización")
        
        self._add_section_title("Localización de la alternativa")
        self._add_subsection_title("01 - Localización de la alternativa")
        
        ubicacion = content.get("ubicacion", {})
        
        table = self.doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = "Ubicación general"
        table.rows[0].cells[1].text = "Ubicación específica"
        for cell in table.rows[0].cells:
            self._set_cell_shading(cell, "0099CC")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        table.rows[1].cells[0].text = f"Región: {ubicacion.get('region', 'Caribe')}"
        table.rows[2].cells[0].text = f"Departamento: {ubicacion.get('departamento', '')}"
        table.rows[3].cells[0].text = f"Municipio: {ubicacion.get('municipio', '')}"
        table.rows[4].cells[0].text = f"Tipo de Agrupación: {ubicacion.get('tipo_agrupacion', 'Urbana')}"
        table.rows[5].cells[0].text = f"Latitud: {ubicacion.get('latitud', '4.123456')}"
        table.rows[6].cells[0].text = f"Longitud: {ubicacion.get('longitud', '-73.987654')}"
        
        self.doc.add_paragraph()
        
        # Factores analizados
        self._add_subsection_title("02 - Factores analizados")
        
        factores = content.get("factores_analizados", [])
        for factor in factores:
            p = self.doc.add_paragraph()
            p.add_run(factor + ",")
    
    def _add_page_14_cadena_valor(self, content: dict):
        """Add Page 14 - Cadena de Valor"""
        self._add_header("Preparación", "Cadena de valor")
        
        self._add_section_title("Cadena de valor de la alternativa")
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Costo total de la alternativa: $ {content.get('costo_total', '')}")
        run.bold = True
        
        obj = content.get("objetivo_especifico", {})
        self._add_subsection_title(f"{obj.get('numero', '1')} - Objetivo específico {obj.get('numero', '1')}   Costo: $ {obj.get('costo', '')}")
        
        p = self.doc.add_paragraph()
        p.add_run(obj.get("descripcion", ""))
        
        self.doc.add_paragraph()
        
        # Producto y actividades table
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = "Producto"
        table.rows[0].cells[1].text = "Actividad y/o Entregable"
        for cell in table.rows[0].cells:
            self._set_cell_shading(cell, "0099CC")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        producto = content.get("producto", {})
        row = table.add_row()
        
        # Left cell - producto
        prod_text = f"{producto.get('codigo', '1.1')} {producto.get('nombre', '')}\n"
        prod_text += f"{producto.get('complemento', '')}\n\n"
        prod_text += f"Complemento:\n"
        prod_text += f"Medido a través de: {producto.get('medido', '')}\n"
        prod_text += f"Cantidad: {producto.get('cantidad', '')}\n"
        prod_text += f"Costo: $ {producto.get('costo', '')}\n"
        prod_text += f"Etapa: {producto.get('etapa', '')}\n"
        prod_text += f"Localización:\n"
        prod_text += f"Número de Personas: {producto.get('num_personas', '')}\n"
        prod_text += f"Acumulativo o no: {producto.get('acumulativo', '')}\n"
        prod_text += f"Población Beneficiaria: {producto.get('poblacion_beneficiaria', '')}"
        row.cells[0].text = prod_text
        
        # Right cell - actividades
        actividades = content.get("actividades", [])
        act_text = ""
        for act in actividades:
            act_text += f"{act.get('codigo', '')} {act.get('descripcion', '')}\n\n"
            act_text += f"Costo: $ {act.get('costo', '')}\n"
            act_text += f"Etapa: {act.get('etapa', '')}\n\n"
        row.cells[1].text = act_text
    
    def _add_page_15_actividades_detalle(self, content: dict):
        """Add Page 15 - Actividades Detalle"""
        self._add_header("Preparación", "Cadena de valor")
        
        actividades = content.get("actividades_periodo", [])
        
        for act in actividades:
            p = self.doc.add_paragraph()
            run = p.add_run(f"Actividad {act.get('codigo', '')} {act.get('codigo', '')} Realizar el pago de los aportes para subsidiar a los usuarios de los servicios públicos domiciliarios de {act.get('nombre', '')} de los estratos uno (1) y dos (2) del municipio.")
            run.bold = True
            run.font.size = Pt(10)
            
            # Periodo table
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            table.rows[0].cells[0].text = "Periodo"
            table.rows[0].cells[1].text = "Servicios domiciliarios"
            for cell in table.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for periodo in self._safe_list(act.get("periodos", [])):
                row = table.add_row()
                row.cells[0].text = str(periodo.get("periodo", "1"))
                row.cells[1].text = periodo.get("valor", "")
            
            row_total = table.add_row()
            row_total.cells[0].text = "Total"
            self._set_cell_shading(row_total.cells[0], "7FC8D8")
            row_total.cells[1].text = act.get("total", "")
            
            self.doc.add_paragraph()
            
            # Total table
            table2 = self.doc.add_table(rows=2, cols=2)
            table2.style = 'Table Grid'
            table2.rows[0].cells[0].text = "Periodo"
            table2.rows[0].cells[1].text = "Total"
            for cell in table2.rows[0].cells:
                self._set_cell_shading(cell, "0099CC")
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            table2.rows[1].cells[0].text = "1"
            table2.rows[1].cells[1].text = act.get("total", "")
            
            self.doc.add_paragraph()
    
    def _add_page_16_riesgos(self, content: dict):
        """Add Page 16 - Riesgos"""
        self._add_header("Preparación", "Riesgos")
        
        self._add_section_title("Análisis de riesgos alternativa")
        self._add_subsection_title("01 - Análisis de riesgo")
        
        # Riesgos table
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        headers = ["Tipo de riesgo", "Descripción del riesgo", "Probabilidad e impacto", "Efectos", "Medidas de mitigación"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            self._set_cell_shading(table.rows[0].cells[i], "0099CC")
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
        
        riesgos = content.get("riesgos", [])
        for riesgo in riesgos:
            row = table.add_row()
            row.cells[0].text = f"{riesgo.get('nivel', '')}\n{riesgo.get('tipo', '')}"
            row.cells[1].text = riesgo.get("descripcion", "")
            row.cells[2].text = f"Probabilidad: {riesgo.get('probabilidad', '')}\n\nImpacto: {riesgo.get('impacto', '')}"
            row.cells[3].text = riesgo.get("efectos", "")
            row.cells[4].text = riesgo.get("mitigacion", "")
    
    def _add_pages_12_16(self, content: dict):
        """Add pages 12-16"""
        # Page 12 - Análisis Técnico
        self._add_page_12_analisis_tecnico(content.get("pagina_12_analisis_tecnico", {}))
        
        self.doc.add_page_break()
        
        # Page 13 - Localización
        self._add_page_13_localizacion(content.get("pagina_13_localizacion", {}))
        
        self.doc.add_page_break()
        
        # Page 14 - Cadena de Valor
        self._add_page_14_cadena_valor(content.get("pagina_14_cadena_valor", {}))
        
        self.doc.add_page_break()
        
        # Page 15 - Actividades Detalle
        self._add_page_15_actividades_detalle(content.get("pagina_15_actividades_detalle", {}))
        
        self.doc.add_page_break()
        
        # Page 16 - Riesgos
        self._add_page_16_riesgos(content.get("pagina_16_riesgos", {}))
    
    def _add_page_17_riesgos_continuacion(self, content: dict):
        """Add Page 17 - Riesgos Continuation"""
        self._add_header("Preparación", "Riesgos")
        
        riesgos = content.get("riesgos_adicionales", [])
        
        for riesgo in riesgos:
            # Risk description
            p = self.doc.add_paragraph()
            p.add_run(riesgo.get("descripcion_actividad", ""))
            p = self.doc.add_paragraph()
            p.add_run(riesgo.get("descripcion_riesgo", ""))
            
            # Add risk table with 4 columns
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            headers = ["Tipo", "Probabilidad/Impacto", "Efectos", "Mitigación"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                self._set_cell_shading(table.rows[0].cells[i], "0099CC")
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(8)
            
            # Data row
            row = table.add_row()
            row.cells[0].text = riesgo.get("tipo", "")
            row.cells[1].text = f"Probabilidad: {riesgo.get('probabilidad', '')}\n\nImpacto: {riesgo.get('impacto', '')}"
            row.cells[2].text = riesgo.get("efectos", "")
            row.cells[3].text = riesgo.get("mitigacion", "")
            
            self.doc.add_paragraph()
    
    def _add_page_18_19_ingresos_beneficios(self, content: dict):
        """Add Pages 18-19 - Ingresos y Beneficios"""
        self._add_header("Preparación", "Ingresos y beneficios")
        
        self._add_section_title("Ingresos y beneficios alternativa")
        self._add_subsection_title("01 - Ingresos y beneficios")
        
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.add_run(content.get("descripcion", ""))
        
        self._add_field("Tipo:", content.get("tipo", "Beneficios"))
        self._add_field("Medido a través de:", content.get("medido", "Número"))
        self._add_field("Bien producido:", content.get("bien_producido", "Otros"))
        self._add_field("Razón Precio Cuenta (RPC):", content.get("razon_precio_cuenta", "0.80"))
        self._add_field("Descripción Cantidad:", content.get("descripcion_cantidad", ""))
        self._add_field("Descripción Valor Unitario:", content.get("descripcion_valor_unitario", ""))
        
        # Periodos table
        tabla_periodos = content.get("tabla_periodos", [])
        if tabla_periodos:
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            headers = ["Periodo", "Cantidad", "Valor unitario", "Valor total"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                self._set_cell_shading(table.rows[0].cells[i], "0099CC")
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for item in tabla_periodos:
                row = table.add_row()
                row.cells[0].text = str(item.get("periodo", ""))
                row.cells[1].text = str(item.get("cantidad", ""))
                row.cells[2].text = str(item.get("valor_unitario", ""))
                row.cells[3].text = str(item.get("valor_total", ""))
        
        self.doc.add_paragraph()
        
        # Totales section
        self._add_subsection_title("02 - Totales")
        
        tabla_totales = content.get("tabla_totales", [])
        if tabla_totales:
            table2 = self.doc.add_table(rows=1, cols=3)
            table2.style = 'Table Grid'
            
            headers = ["Periodo", "Total beneficios", "Total"]
            for i, h in enumerate(headers):
                table2.rows[0].cells[i].text = h
                self._set_cell_shading(table2.rows[0].cells[i], "0099CC")
                for run in table2.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for item in tabla_totales:
                row = table2.add_row()
                row.cells[0].text = str(item.get("periodo", ""))
                row.cells[1].text = str(item.get("total_beneficios", ""))
                row.cells[2].text = str(item.get("total", ""))
    
    def _add_page_20_flujo_economico(self, content: dict):
        """Add Page 20 - Flujo Económico"""
        self._add_header("Evaluación", "Flujo Económico")
        
        p = self.doc.add_paragraph()
        run = p.add_run(content.get("alternativa", "Alternativa 1"))
        run.bold = True
        
        self._add_section_title("Flujo")
        self._add_subsection_title("01 - Flujo Económico")
        
        flujo = content.get("flujo", [])
        if flujo:
            table = self.doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            
            headers = ["P", "Beneficios e ingresos (+)", "Créditos(+)", "Costos de preinversión (-)", 
                      "Costos de inversión (-)", "Costos de operación (-)", "Amortización (-)", 
                      "Intereses de los créditos (-)", "Flujo Neto"]
            
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                self._set_cell_shading(table.rows[0].cells[i], "0099CC")
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(7)
            
            for item in flujo:
                row = table.add_row()
                row.cells[0].text = str(item.get("p", ""))
                row.cells[1].text = str(item.get("beneficios", ""))
                row.cells[2].text = str(item.get("creditos", ""))
                row.cells[3].text = str(item.get("costos_preinversion", ""))
                row.cells[4].text = str(item.get("costos_inversion", ""))
                row.cells[5].text = str(item.get("costos_operacion", ""))
                row.cells[6].text = str(item.get("amortizacion", ""))
                row.cells[7].text = str(item.get("intereses", ""))
                row.cells[8].text = str(item.get("flujo_neto", ""))
    
    def _add_page_21_indicadores_decision(self, content: dict):
        """Add Page 21 - Indicadores y Decisión"""
        self._add_header("Evaluación", "Indicadores y decisión")
        
        self._add_section_title("Indicadores y decisión")
        
        # Evaluación económica
        self._add_subsection_title("01 - Evaluación económica")
        
        eval_eco = content.get("evaluacion_economica", {})
        
        # Complex header table
        table = self.doc.add_table(rows=3, cols=6)
        table.style = 'Table Grid'
        
        # Header row 1
        table.rows[0].cells[0].text = "Indicadores de rentabilidad"
        self._set_cell_shading(table.rows[0].cells[0], "0099CC")
        table.rows[0].cells[3].text = "Indicadores de costo-eficiencia"
        self._set_cell_shading(table.rows[0].cells[3], "0099CC")
        table.rows[0].cells[4].text = "Indicadores de costo mínimo"
        self._set_cell_shading(table.rows[0].cells[4], "0099CC")
        
        # Header row 2
        headers2 = ["Valor Presente Neto (VPN)", "Tasa Interna de Retorno (TIR)", 
                   "Relación Costo Beneficio (RCB)", "Costo por beneficiario",
                   "Valor presente de los costos", "Costo Anual Equivalente (CAE)"]
        for i, h in enumerate(headers2):
            table.rows[1].cells[i].text = h
            self._set_cell_shading(table.rows[1].cells[i], "7FC8D8")
            for run in table.rows[1].cells[i].paragraphs[0].runs:
                run.font.size = Pt(8)
        
        # Values row
        table.rows[2].cells[0].text = eval_eco.get("vpn", "")
        table.rows[2].cells[1].text = eval_eco.get("tir", "")
        table.rows[2].cells[2].text = eval_eco.get("rcb", "")
        table.rows[2].cells[3].text = eval_eco.get("costo_beneficiario", "")
        table.rows[2].cells[4].text = eval_eco.get("valor_presente_costos", "")
        table.rows[2].cells[5].text = eval_eco.get("cae", "")
        
        self.doc.add_paragraph()
        
        # Costo por capacidad
        self._add_subsection_title("Costo por capacidad")
        
        costo_cap = content.get("costo_capacidad", {})
        
        table2 = self.doc.add_table(rows=2, cols=2)
        table2.style = 'Table Grid'
        table2.rows[0].cells[0].text = "Producto"
        table2.rows[0].cells[1].text = "Costo unitario (valor presente)"
        for cell in table2.rows[0].cells:
            self._set_cell_shading(cell, "0099CC")
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        table2.rows[1].cells[0].text = costo_cap.get("producto", "")
        table2.rows[1].cells[1].text = costo_cap.get("costo_unitario", "")
        
        self.doc.add_paragraph()
        
        # Decisión
        self._add_subsection_title("03 - Decisión")
        
        decision = content.get("decision", {})
        self._add_field("Alternativa", decision.get("alternativa", ""))
        
        self.doc.add_paragraph()
        
        # Alcance
        self._add_subsection_title("04 - Alcance")
        
        p = self.doc.add_paragraph()
        p.add_run(content.get("alcance", ""))
    
    def _add_pages_17_21(self, content: dict):
        """Add pages 17-21"""
        # Page 17 - Riesgos Continuation
        self._add_page_17_riesgos_continuacion(content.get("pagina_17_riesgos_continuacion", {}))
        
        self.doc.add_page_break()
        
        # Pages 18-19 - Ingresos y Beneficios
        self._add_page_18_19_ingresos_beneficios(content.get("pagina_18_19_ingresos_beneficios", {}))
        
        self.doc.add_page_break()
        
        # Page 20 - Flujo Económico
        self._add_page_20_flujo_economico(content.get("pagina_20_flujo_economico", {}))
        
        self.doc.add_page_break()
        
        # Page 21 - Indicadores y Decisión
        self._add_page_21_indicadores_decision(content.get("pagina_21_indicadores_decision", {}))
    
    def _add_page_22_indicadores_producto(self, content: dict):
        """Add Page 22 - Indicadores de Producto"""
        self._add_header("Programación", "Indicadores de producto")
        
        self._add_section_title("Indicadores de producto")
        
        # Objetivo
        objetivo = content.get("objetivo", {})
        self._add_subsection_title(f"01 - Objetivo {objetivo.get('numero', '1')}")
        
        p = self.doc.add_paragraph()
        p.add_run(f"{objetivo.get('numero', '1')}. {objetivo.get('descripcion', '')}")
        
        self.doc.add_paragraph()
        
        # Producto
        producto = content.get("producto", {})
        p = self.doc.add_paragraph()
        run = p.add_run("Producto")
        run.font.color.rgb = RGBColor(0, 153, 204)
        
        p = self.doc.add_paragraph()
        p.add_run(f"{producto.get('codigo', '1.1')}. {producto.get('nombre', '')}     {producto.get('complemento', '')}")
        
        self.doc.add_paragraph()
        
        # Indicador
        indicador = content.get("indicador", {})
        p = self.doc.add_paragraph()
        run = p.add_run("Indicador")
        run.font.color.rgb = RGBColor(0, 153, 204)
        
        p = self.doc.add_paragraph()
        p.add_run(f"{indicador.get('codigo', '1.1.1')} {indicador.get('nombre', '')}")
        
        self._add_field("Medido a través de:", indicador.get("medido", ""))
        self._add_field("Meta total:", indicador.get("meta_total", ""))
        self._add_field("Fórmula:", indicador.get("formula", ""))
        self._add_field("Es acumulativo:", indicador.get("es_acumulativo", ""))
        self._add_field("Es Principal:", indicador.get("es_principal", ""))
        self._add_field("Tipo de Fuente:", indicador.get("tipo_fuente", ""))
        self._add_field("Fuente de Verificación:", indicador.get("fuente_verificacion", ""))
        
        self.doc.add_paragraph()
        
        # Programación de indicadores
        p = self.doc.add_paragraph()
        run = p.add_run("Programación de indicadores")
        run.font.color.rgb = RGBColor(0, 153, 204)
        
        programacion = content.get("programacion_indicadores", [])
        if programacion:
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            headers = ["Periodo", "Meta por periodo", "Periodo", "Meta por periodo"]
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                self._set_cell_shading(table.rows[0].cells[i], "0099CC")
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            for item in programacion:
                row = table.add_row()
                row.cells[0].text = str(item.get("periodo", ""))
                row.cells[1].text = str(item.get("meta", ""))
    
    def _add_page_23_regionalizacion(self, content: dict):
        """Add Page 23 - Regionalización"""
        self._add_header("Programación", "Regionalización")
        
        self._add_section_title("Regionalización")
        
        p = self.doc.add_paragraph()
        run = p.add_run("Producto: ")
        run.bold = True
        p.add_run(content.get("producto", ""))
        
        self.doc.add_paragraph()
        
        # Ubicación table
        ubicacion = content.get("ubicacion", {})
        table = self.doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        
        headers = ["Región", "Departamento", "Municipio", "Tipo de Agrupación", "Agrupación"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            self._set_cell_shading(table.rows[0].cells[i], "0099CC")
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        table.rows[1].cells[0].text = ubicacion.get("region", "Caribe")
        table.rows[1].cells[1].text = ubicacion.get("departamento", "")
        table.rows[1].cells[2].text = ubicacion.get("municipio", "")
        table.rows[1].cells[3].text = ubicacion.get("tipo_agrupacion", "")
        table.rows[1].cells[4].text = ubicacion.get("agrupacion", "")
        
        self.doc.add_paragraph()
        
        # Costos table
        tabla_costos = content.get("tabla_costos", [])
        if tabla_costos:
            table2 = self.doc.add_table(rows=1, cols=6)
            table2.style = 'Table Grid'
            
            headers2 = ["Periodo", "Costo Total", "Costo Regionalizado", "Meta Total", "Meta Regionalizada", "Beneficiarios"]
            for i, h in enumerate(headers2):
                table2.rows[0].cells[i].text = h
                self._set_cell_shading(table2.rows[0].cells[i], "0099CC")
                for run in table2.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(8)
            
            for item in tabla_costos:
                row = table2.add_row()
                row.cells[0].text = str(item.get("periodo", ""))
                row.cells[1].text = str(item.get("costo_total", ""))
                row.cells[2].text = str(item.get("costo_regionalizado", ""))
                row.cells[3].text = str(item.get("meta_total", ""))
                row.cells[4].text = str(item.get("meta_regionalizada", ""))
                row.cells[5].text = str(item.get("beneficiarios", ""))
    
    def _add_page_24_focalizacion(self, content: dict):
        """Add Page 24 - Focalización"""
        self._add_header("Programación", "Focalización")
        
        self._add_section_title("Focalización")
        
        # Focalización table (empty in template)
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ["Política", "Categoría", "SubCategoría", "Valor"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            self._set_cell_shading(table.rows[0].cells[i], "0099CC")
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Add empty rows if data exists
        tabla_focalizacion = content.get("tabla_focalizacion", [])
        for item in tabla_focalizacion:
            row = table.add_row()
            row.cells[0].text = str(item.get("politica", ""))
            row.cells[1].text = str(item.get("categoria", ""))
            row.cells[2].text = str(item.get("subcategoria", ""))
            row.cells[3].text = str(item.get("valor", ""))
    
    def _add_pages_22_24(self, content: dict):
        """Add pages 22-24"""
        # Page 22 - Indicadores de Producto
        self._add_page_22_indicadores_producto(content.get("pagina_22_indicadores_producto", {}))
        
        self.doc.add_page_break()
        
        # Page 23 - Regionalización
        self._add_page_23_regionalizacion(content.get("pagina_23_regionalizacion", {}))
        
        self.doc.add_page_break()
        
        # Page 24 - Focalización
        self._add_page_24_focalizacion(content.get("pagina_24_focalizacion", {}))
    
    def _set_cell_shading(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        tcPr.append(shading)
    
    def _save_document(self, data):
        """Save the document to file as PDF"""
        from docx2pdf import convert
        import platform
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Handle case where data is not a dict
        if isinstance(data, dict):
            municipio = data.get("municipio", "documento")
        else:
            municipio = "documento"
        
        # Remove invalid filename characters
        municipio = re.sub(r'[^a-zA-Z0-9_\-]', '', municipio.replace(" ", "_"))
        if not municipio:
            municipio = "documento"
        
        # First save as DOCX
        docx_filename = f"MGA_Subsidios_{municipio}_{timestamp}.docx"
        docx_filepath = os.path.join(self.output_dir, docx_filename)
        self.doc.save(docx_filepath)
        
        # Convert to PDF
        pdf_filename = f"MGA_Subsidios_{municipio}_{timestamp}.pdf"
        pdf_filepath = os.path.join(self.output_dir, pdf_filename)
        
        try:
            # Initialize COM for Windows (required for docx2pdf in threaded environments)
            if platform.system() == 'Windows':
                import pythoncom
                pythoncom.CoInitialize()
            
            convert(docx_filepath, pdf_filepath)
            
            # Uninitialize COM
            if platform.system() == 'Windows':
                pythoncom.CoUninitialize()
            
            # Remove temporary DOCX file
            os.remove(docx_filepath)
            return pdf_filepath
        except Exception as e:
            # If PDF conversion fails, return DOCX
            print(f"PDF conversion failed: {e}, returning DOCX")
            return docx_filepath

